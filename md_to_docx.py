#!/usr/bin/env python3

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# Full-line markdown image: ![alt](path)
IMG_PATTERN = re.compile(r"!\[(.*?)\]\(([^)]+)\)")

# Markdown ATX heading: #, ##, ..., ###### at start of line
HEADING_PATTERN = re.compile(r"^(#{1,6})\s+(.+)$")


def _add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    """
    Оборачиваем заголовок в bookmark, чтобы на него можно было сослаться
    из оглавления.
    """
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)

    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))

    p = paragraph._p
    # bookmarkStart в начало параграфа, bookmarkEnd — в конец
    p.insert(0, start)
    p.append(end)


def _add_toc_entry(doc: Document, text: str, bookmark_name: str, level: int) -> None:
    """
    Добавляем строку оглавления как гиперссылку на bookmark соответствующего заголовка.
    """
    p = doc.add_paragraph()
    if level > 1:
        p.paragraph_format.left_indent = Inches(0.25 * (level - 1))

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    hyperlink.set(qn("w:history"), "1")

    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)
    r.append(r_pr)

    t = OxmlElement("w:t")
    t.text = text
    r.append(t)

    hyperlink.append(r)
    p._p.append(hyperlink)


def md_to_docx(md_path: str, docx_path: str | None = None) -> Path:
    """
    Convert a Markdown file to DOCX, replacing full-line image markdown
    like `![alt](path)` with actual images.
    """

    md_path_obj = Path(md_path).expanduser().resolve()
    if docx_path is None:
        docx_path_obj = md_path_obj.with_suffix(".docx")
    else:
        docx_path_obj = Path(docx_path).expanduser().resolve()

    if not md_path_obj.exists():
        raise FileNotFoundError(f"Input markdown not found: {md_path_obj}")

    text = md_path_obj.read_text(encoding="utf-8")
    lines = text.splitlines()

    # Первый проход: собираем все заголовки, чтобы построить оглавление.
    all_headings: list[tuple[int, str]] = []
    for line in lines:
        m = HEADING_PATTERN.match(line.strip())
        if m:
            hashes, title = m.groups()
            level = min(len(hashes), 4)
            all_headings.append((level, title.strip()))

    doc = Document()
    assets_root = md_path_obj.parent

    in_toc = False
    heading_index = 0  # сквозная нумерация заголовков для bookmark'ов

    for line in lines:
        stripped = line.strip()

        # Headings: "# Title" -> Word "Heading 1", etc.
        heading_match = HEADING_PATTERN.match(stripped)
        if heading_match:
            hashes, title = heading_match.groups()
            level = min(len(hashes), 4)  # map 1–4, 5/6 тоже в Heading 4
            style = f"Heading {level}"
            heading_index += 1
            bookmark_name = f"h{heading_index}"

            title_clean = title.strip()
            # Считаем, что после "Оглавление" в markdown идёт ручное оглавление,
            # его пропускаем и вместо него строим своё, кликабельное.
            if title_clean.lower() == "оглавление":
                p = doc.add_paragraph(title_clean, style=style)
                _add_bookmark(p, bookmark_name, heading_index)
                in_toc = True

                # Добавляем строки оглавления для всех остальных заголовков
                for idx, (lvl, t) in enumerate(all_headings, start=1):
                    if t.lower() == "оглавление":
                        continue
                    _add_toc_entry(doc, t, f"h{idx}", lvl)
            else:
                p = doc.add_paragraph(title_clean, style=style)
                _add_bookmark(p, bookmark_name, heading_index)
            continue

        # Пропускаем оригинальное markdown-оглавление до разделителя '---'
        if in_toc:
            if stripped == "---":
                in_toc = False
            continue

        img_match = IMG_PATTERN.fullmatch(stripped)
        if img_match:
            alt, rel_path = img_match.groups()
            img_path = (assets_root / rel_path).resolve()

            if img_path.exists():
                if alt:
                    doc.add_paragraph(alt)
                try:
                    doc.add_picture(str(img_path), width=Inches(5))
                except Exception as exc:
                    doc.add_paragraph(f"[image error: {alt}] {rel_path} ({exc})")
            else:
                # If the image file is missing, keep the original markdown line
                doc.add_paragraph(line)
        else:
            doc.add_paragraph(line)

    doc.save(docx_path_obj)
    return docx_path_obj


def main(argv: list[str]) -> int:
    if len(argv) < 2:
        print("Usage: md_to_docx.py INPUT_MD [OUTPUT_DOCX]", file=sys.stderr)
        return 1

    md_path = argv[1]
    docx_path = argv[2] if len(argv) > 2 else None

    out = md_to_docx(md_path, docx_path)
    print(out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
