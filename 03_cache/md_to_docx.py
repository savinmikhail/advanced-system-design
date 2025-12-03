#!/usr/bin/env python3

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches


IMG_PATTERN = re.compile(r"!\\[(.*?)\\]\\(([^)]+)\\)")


def md_to_docx(md_path: str, docx_path: str | None = None) -> Path:
    """
    Convert a Markdown file to DOCX, replacing full-line image markdown
    like `![alt](path)` with actual images.
    """

    md_path_obj = Path(md_path).expanduser().resolve()
    if docx_path is None:
        docx_path_obj = md_path_obj.with_suffix(".docx")
    else:
        docx_path_obj = Path(docx_path).expanduser().resolve()

    if not md_path_obj.exists():
        raise FileNotFoundError(f"Input markdown not found: {md_path_obj}")

    text = md_path_obj.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    assets_root = md_path_obj.parent

    for line in lines:
        stripped = line.strip()
        match = IMG_PATTERN.fullmatch(stripped)

        if match:
            alt, rel_path = match.groups()
            img_path = (assets_root / rel_path).resolve()

            if img_path.exists():
                if alt:
                    doc.add_paragraph(alt)
                try:
                    doc.add_picture(str(img_path), width=Inches(5))
                except Exception as exc:
                    doc.add_paragraph(f"[image error: {alt}] {rel_path} ({exc})")
            else:
                # If the image file is missing, keep the original markdown line
                doc.add_paragraph(line)
        else:
            doc.add_paragraph(line)

    doc.save(docx_path_obj)
    return docx_path_obj


def main(argv: list[str]) -> int:
    if len(argv) < 2:
        print("Usage: md_to_docx.py INPUT_MD [OUTPUT_DOCX]", file=sys.stderr)
        return 1

    md_path = argv[1]
    docx_path = argv[2] if len(argv) > 2 else None

    out = md_to_docx(md_path, docx_path)
    print(out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))

